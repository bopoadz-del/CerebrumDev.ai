"""Ingestion: parse formats, chunk with provenance, embed, persist.

Supported: PDF, DOCX, TXT, Markdown, CSV, XLSX. Trusted tenant/project/layer
must be supplied by the caller (server-resolved) — never from model output.
"""

from __future__ import annotations

import csv
import io
import uuid
from dataclasses import dataclass
from typing import Any, Dict, List, Optional

from sqlalchemy.orm import Session

from app.steward.embedding_guard import assert_fingerprint_compatible
from app.steward.embeddings import get_embedder
from app.steward.models import Document, DocumentChunk

CHUNK_SIZE = 800
CHUNK_OVERLAP = 120
ROWS_PER_CHUNK = 25


@dataclass
class ParsedUnit:
    text: str
    page: Optional[int] = None
    sheet: Optional[str] = None
    row_start: Optional[int] = None
    row_end: Optional[int] = None
    rows: Optional[List[Dict[str, Any]]] = None
    unit_type: str = "text"


def detect_type(filename: str, content_type: Optional[str] = None) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith((".md", ".markdown")):
        return "markdown"
    if name.endswith(".csv"):
        return "csv"
    if name.endswith((".xlsx", ".xlsm")):
        return "xlsx"
    if name.endswith((".txt", ".text")):
        return "txt"
    if content_type:
        if "pdf" in content_type:
            return "pdf"
        if "csv" in content_type:
            return "csv"
        if "sheet" in content_type or "excel" in content_type:
            return "xlsx"
        if "word" in content_type:
            return "docx"
    return "txt"


def _parse_text(data: bytes) -> List[ParsedUnit]:
    text = data.decode("utf-8", errors="replace")
    return [ParsedUnit(text=text, unit_type="text")]


def _parse_pdf(data: bytes) -> List[ParsedUnit]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    units: List[ParsedUnit] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            units.append(ParsedUnit(text=text, page=i, unit_type="text"))
    return units


def _parse_docx(data: bytes) -> List[ParsedUnit]:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return [ParsedUnit(text="\n".join(parts), unit_type="text")] if parts else []


def _rows_to_text(header: List[str], rows: List[Dict[str, Any]]) -> str:
    lines = [", ".join(header)]
    for row in rows:
        lines.append(", ".join(f"{k}={row.get(k, '')}" for k in header))
    return "\n".join(lines)


def _rows_into_units(
    header: List[str], rows: List[Dict[str, Any]], sheet: str
) -> List[ParsedUnit]:
    units: List[ParsedUnit] = []
    for start in range(0, len(rows), ROWS_PER_CHUNK):
        block = rows[start : start + ROWS_PER_CHUNK]
        units.append(
            ParsedUnit(
                text=_rows_to_text(header, block),
                sheet=sheet,
                row_start=start + 2,
                row_end=start + 1 + len(block),
                rows=block,
                unit_type="table",
            )
        )
    return units


def _parse_csv(data: bytes, sheet: str = "sheet1") -> List[ParsedUnit]:
    text = data.decode("utf-8-sig", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    header = list(reader.fieldnames or [])
    all_rows = [dict(r) for r in reader]
    return _rows_into_units(header, all_rows, sheet)


def _parse_xlsx(data: bytes) -> List[ParsedUnit]:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    units: List[ParsedUnit] = []
    for ws in wb.worksheets:
        rows_iter = ws.iter_rows(values_only=True)
        try:
            header_row = next(rows_iter)
        except StopIteration:
            continue
        header = [str(h) if h is not None else f"col{i}" for i, h in enumerate(header_row)]
        data_rows: List[Dict[str, Any]] = []
        for values in rows_iter:
            if values is None or all(v is None for v in values):
                continue
            data_rows.append(
                {header[i]: values[i] for i in range(min(len(header), len(values)))}
            )
        units.extend(_rows_into_units(header, data_rows, ws.title))
    return units


_PARSERS = {
    "txt": _parse_text,
    "markdown": _parse_text,
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "csv": _parse_csv,
    "xlsx": _parse_xlsx,
}


def parse_bytes(
    data: bytes, filename: str, content_type: Optional[str] = None
) -> List[ParsedUnit]:
    kind = detect_type(filename, content_type)
    parser = _PARSERS.get(kind, _parse_text)
    return parser(data)


def _chunk_text(text: str) -> List[str]:
    text = text.strip()
    if len(text) <= CHUNK_SIZE:
        return [text] if text else []
    chunks: List[str] = []
    start = 0
    while start < len(text):
        end = min(start + CHUNK_SIZE, len(text))
        window = text[start:end]
        if end < len(text):
            for sep in ("\n\n", "\n", ". "):
                idx = window.rfind(sep)
                if idx > CHUNK_SIZE // 2:
                    end = start + idx + len(sep)
                    window = text[start:end]
                    break
        chunks.append(window.strip())
        if end >= len(text):
            break
        start = max(end - CHUNK_OVERLAP, start + 1)
    return [c for c in chunks if c]


def ingest_bytes(
    session: Session,
    *,
    tenant_id: str,
    project_id: str,
    knowledge_layer: int,
    filename: str,
    data: bytes,
    content_type: Optional[str] = None,
    rag_pack_id: Optional[str] = None,
    collection_id: Optional[str] = None,
    source_record_id: Optional[str] = None,
    authority_kind: str = "estate_owned",
    cite_as_policy: bool = True,
    document_id: Optional[str] = None,
    meta: Optional[Dict[str, Any]] = None,
    prior_fingerprint: Optional[str] = None,
) -> Document:
    if not tenant_id or not project_id:
        raise ValueError("tenant_id and project_id are required")
    if knowledge_layer not in (1, 2):
        raise ValueError("knowledge_layer must be 1 or 2")
    if authority_kind in {"synthetic", "evaluation_only"} and cite_as_policy:
        raise ValueError("synthetic/evaluation content cannot be cited as policy")

    embedder = get_embedder()
    assert_fingerprint_compatible(prior_fingerprint)

    units = parse_bytes(data, filename, content_type)
    doc_id = document_id or uuid.uuid4().hex
    doc = Document(
        id=doc_id,
        tenant_id=tenant_id,
        project_id=project_id,
        knowledge_layer=knowledge_layer,
        filename=filename,
        content_type=content_type or detect_type(filename),
        rag_pack_id=rag_pack_id,
        collection_id=collection_id,
        source_record_id=source_record_id,
        status="processing",
        authority_kind=authority_kind,
        embedding_fingerprint=embedder.fingerprint,
        meta=meta or {},
    )
    session.add(doc)

    ordinal = 0
    page_max = 0
    for unit in units:
        pieces = _chunk_text(unit.text)
        vectors = embedder.embed_batch(pieces) if pieces else []
        for text, vector in zip(pieces, vectors):
            chunk = DocumentChunk(
                id=uuid.uuid4().hex,
                tenant_id=tenant_id,
                project_id=project_id,
                knowledge_layer=knowledge_layer,
                document_id=doc_id,
                chunk_ordinal=ordinal,
                text=text,
                embedding=vector,
                source_filename=filename,
                page=unit.page,
                sheet=unit.sheet,
                row_start=unit.row_start,
                row_end=unit.row_end,
                rag_pack_id=rag_pack_id,
                authority_kind=authority_kind,
                cite_as_policy=cite_as_policy,
                meta={"unit_type": unit.unit_type, **({"rows": unit.rows} if unit.rows else {})},
            )
            session.add(chunk)
            ordinal += 1
        if unit.page:
            page_max = max(page_max, unit.page)

    doc.chunk_count = ordinal
    doc.page_count = page_max or None
    doc.status = "ready" if ordinal else "empty"
    session.flush()
    return doc
