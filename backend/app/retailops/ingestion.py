"""Ingestion: parse pilot formats, chunk with provenance, embed, and persist.

Supported formats: PDF, DOCX, TXT, Markdown, CSV, XLSX. CSV/XLSX preserve
sheet + row provenance and also expose a normalized structured-row payload
(``chunk.meta['rows']``) that structured actions (e.g. inventory analysis)
consume, while still producing textual chunks for retrieval.
"""

from __future__ import annotations

import csv
import io
import uuid
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

from sqlalchemy.orm import Session

from app.retailops.embeddings import get_embedder
from app.retailops.models import Document, DocumentChunk

CHUNK_SIZE = 800
CHUNK_OVERLAP = 120
ROWS_PER_CHUNK = 25


@dataclass
class ParsedUnit:
    text: str
    page: Optional[int] = None
    sheet: Optional[str] = None
    row_start: Optional[int] = None
    row_end: Optional[int] = None
    rows: Optional[List[Dict[str, Any]]] = None
    unit_type: str = "text"


# ---------------------------------------------------------------------------
# Format detection
# ---------------------------------------------------------------------------

def detect_type(filename: str, content_type: Optional[str] = None) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith((".md", ".markdown")):
        return "markdown"
    if name.endswith(".csv"):
        return "csv"
    if name.endswith((".xlsx", ".xlsm")):
        return "xlsx"
    if name.endswith((".txt", ".text")):
        return "txt"
    if content_type:
        if "pdf" in content_type:
            return "pdf"
        if "csv" in content_type:
            return "csv"
        if "sheet" in content_type or "excel" in content_type:
            return "xlsx"
        if "word" in content_type:
            return "docx"
    return "txt"


# ---------------------------------------------------------------------------
# Parsers
# ---------------------------------------------------------------------------

def _parse_text(data: bytes) -> List[ParsedUnit]:
    text = data.decode("utf-8", errors="replace")
    return [ParsedUnit(text=text, unit_type="text")]


def _parse_pdf(data: bytes) -> List[ParsedUnit]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    units: List[ParsedUnit] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            units.append(ParsedUnit(text=text, page=i, unit_type="text"))
    return units


def _parse_docx(data: bytes) -> List[ParsedUnit]:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return [ParsedUnit(text="\n".join(parts), unit_type="text")] if parts else []


def _rows_to_text(header: List[str], rows: List[Dict[str, Any]]) -> str:
    lines = [", ".join(header)]
    for row in rows:
        lines.append(", ".join(f"{k}={row.get(k, '')}" for k in header))
    return "\n".join(lines)


def _parse_csv(data: bytes, sheet: str = "sheet1") -> List[ParsedUnit]:
    text = data.decode("utf-8-sig", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    header = reader.fieldnames or []
    all_rows = [dict(r) for r in reader]
    return _rows_into_units(header, all_rows, sheet)


def _parse_xlsx(data: bytes) -> List[ParsedUnit]:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    units: List[ParsedUnit] = []
    for ws in wb.worksheets:
        rows_iter = ws.iter_rows(values_only=True)
        try:
            header_row = next(rows_iter)
        except StopIteration:
            continue
        header = [str(h) if h is not None else f"col{i}" for i, h in enumerate(header_row)]
        data_rows: List[Dict[str, Any]] = []
        for values in rows_iter:
            if values is None or all(v is None for v in values):
                continue
            data_rows.append({header[i]: values[i] for i in range(min(len(header), len(values)))})
        units.extend(_rows_into_units(header, data_rows, ws.title))
    return units


def _rows_into_units(
    header: List[str], rows: List[Dict[str, Any]], sheet: str
) -> List[ParsedUnit]:
    units: List[ParsedUnit] = []
    for start in range(0, len(rows), ROWS_PER_CHUNK):
        block = rows[start : start + ROWS_PER_CHUNK]
        units.append(
            ParsedUnit(
                text=_rows_to_text(header, block),
                sheet=sheet,
                row_start=start + 2,  # +1 for header, +1 for 1-based
                row_end=start + 1 + len(block),
                rows=block,
                unit_type="table",
            )
        )
    return units


_PARSERS = {
    "txt": _parse_text,
    "markdown": _parse_text,
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "csv": _parse_csv,
    "xlsx": _parse_xlsx,
}


def parse_bytes(data: bytes, filename: str, content_type: Optional[str] = None) -> List[ParsedUnit]:
    kind = detect_type(filename, content_type)
    parser = _PARSERS.get(kind, _parse_text)
    return parser(data)


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def _chunk_text(text: str) -> List[str]:
    text = text.strip()
    if len(text) <= CHUNK_SIZE:
        return [text] if text else []
    chunks: List[str] = []
    start = 0
    while start < len(text):
        end = min(start + CHUNK_SIZE, len(text))
        # try to break on a paragraph/sentence boundary
        window = text[start:end]
        if end < len(text):
            for sep in ("\n\n", "\n", ". "):
                idx = window.rfind(sep)
                if idx > CHUNK_SIZE // 2:
                    end = start + idx + len(sep)
                    window = text[start:end]
                    break
        chunks.append(window.strip())
        start = max(end - CHUNK_OVERLAP, end) if end == len(text) else end - CHUNK_OVERLAP
    return [c for c in chunks if c]


def units_to_chunks(units: List[ParsedUnit]) -> List[ParsedUnit]:
    out: List[ParsedUnit] = []
    for unit in units:
        if unit.unit_type == "table":
            out.append(unit)
            continue
        for piece in _chunk_text(unit.text):
            out.append(
                ParsedUnit(
                    text=piece,
                    page=unit.page,
                    sheet=unit.sheet,
                    unit_type="text",
                )
            )
    return out


# ---------------------------------------------------------------------------
# Persistence
# ---------------------------------------------------------------------------

def ingest_document(
    session: Session,
    *,
    tenant_id: str,
    project_id: str,
    filename: str,
    data: bytes,
    content_type: Optional[str] = None,
    source_uri: Optional[str] = None,
    document_id: Optional[str] = None,
) -> Document:
    """Parse, chunk, embed and persist a document + its chunks (production path)."""
    kind = detect_type(filename, content_type)
    document_id = document_id or f"doc_{uuid.uuid4().hex[:12]}"
    doc = Document(
        id=document_id,
        tenant_id=tenant_id,
        project_id=project_id,
        filename=filename,
        content_type=content_type or kind,
        source_uri=source_uri,
        status="processing",
        meta={"format": kind},
    )
    session.add(doc)
    session.flush()

    try:
        units = parse_bytes(data, filename, content_type)
        chunks = units_to_chunks(units)
        embedder = get_embedder()
        texts = [c.text for c in chunks]
        vectors = embedder.embed_batch(texts) if texts else []
        pages = {u.page for u in units if u.page}
        for ordinal, (unit, vector) in enumerate(zip(chunks, vectors)):
            meta: Dict[str, Any] = {"unit_type": unit.unit_type}
            if unit.rows is not None:
                meta["rows"] = unit.rows
            session.add(
                DocumentChunk(
                    id=f"ch_{uuid.uuid4().hex[:16]}",
                    tenant_id=tenant_id,
                    project_id=project_id,
                    document_id=document_id,
                    chunk_ordinal=ordinal,
                    text=unit.text,
                    embedding=vector,
                    source_filename=filename,
                    page=unit.page,
                    sheet=unit.sheet,
                    row_start=unit.row_start,
                    row_end=unit.row_end,
                    meta=meta,
                )
            )
        doc.status = "processed"
        doc.chunk_count = len(chunks)
        doc.page_count = len(pages) or None
    except Exception as exc:  # noqa: BLE001
        doc.status = "failed"
        doc.error = str(exc)[:500]
        session.flush()
        raise
    session.flush()
    return doc
